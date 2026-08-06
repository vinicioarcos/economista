from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docx" / "formulario-datos-portafolio.docx"
PUBLIC = ROOT / "public" / "downloads" / "formulario-datos-portafolio.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)
PUBLIC.parent.mkdir(parents=True, exist_ok=True)

GREEN = "1B4D3E"
LIME = "D7F05B"
INK = "17211B"
MUTED = "647067"
LINE = "D7D8CF"
PALE = "F4F1E8"

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.78)
section.right_margin = Inches(1)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

def set_font(run, size=11, bold=False, color=INK, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for style_name, size, before, after in [("Title", 25, 0, 5), ("Subtitle", 12, 0, 18), ("Heading 1", 16, 18, 10), ("Heading 2", 13, 14, 7), ("Heading 3", 12, 10, 5)]:
    s = styles[style_name]
    s.font.name = "Calibri"
    s.font.size = Pt(size)
    s.font.bold = style_name != "Subtitle"
    s.font.color.rgb = RGBColor.from_string(GREEN if style_name != "Subtitle" else MUTED)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def header_footer(current_section):
    hp = current_section.header.paragraphs[0]
    hp.text = "PORTAFOLIO PROFESIONAL DE ECONOMÍA  |  FORMULARIO DE INFORMACIÓN"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    set_font(hp.runs[0], 7.5, True, MUTED)
    fp = current_section.footer.paragraphs[0]
    fp.text = "Complete solo información verificable. No incluya contraseñas, cédula ni domicilio exacto."
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_after = Pt(0)
    set_font(fp.runs[0], 7.5, False, MUTED)

header_footer(section)

def field(label, hint="", lines=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    set_font(r, 10, True, GREEN)
    if hint:
        r = p.add_run(f"  {hint}")
        set_font(r, 8.5, False, MUTED)
    for _ in range(lines):
        p = doc.add_paragraph("________________________________________________________________________________")
        p.paragraph_format.space_after = Pt(3)
        set_font(p.runs[0], 10, False, LINE)

def note(text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    shade(table.cell(0, 0), PALE)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, 9, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def page_break():
    # El formulario fluye de manera continua. Los estilos de encabezado usan
    # keep_with_next para evitar títulos huérfanos sin insertar páginas vacías.
    return None

def repeated_table(headers, widths, rows=3):
    table = doc.add_table(rows=1 + rows, cols=len(headers))
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        shade(table.cell(0, i), GREEN)
        p = table.cell(0, i).paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_font(r, 8.5, True, "FFFFFF")
    for row in table.rows[1:]:
        for cell in row.cells:
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(20)
            r = p.add_run("Escriba aquí")
            set_font(r, 8, False, MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# Page 1 - customer-pack style opening
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(5)
r = p.add_run("FORMULARIO DE INFORMACIÓN")
set_font(r, 9, True, GREEN)
p = doc.add_paragraph(style="Title")
p.add_run("Portafolio profesional\nde Economía")
p = doc.add_paragraph(style="Subtitle")
p.add_run("Datos para construir la web, la hoja de vida y la marca personal")
note("Propósito: recopilar evidencia suficiente para que la IA prepare un portafolio veraz. Si un dato no aplica, escriba NO APLICA. Si todavía no tiene un enlace, escriba PENDIENTE.")

doc.add_heading("1. Identidad profesional", level=1)
field("Nombre completo")
field("Nombre preferido para la web")
field("Ciudad y país")
field("Titular profesional", "Máximo 12 palabras. Ej.: Estudiante de Economía | Datos y políticas públicas")
field("Correo autorizado para publicación")
field("Teléfono autorizado para publicación", "Opcional")
field("Tipo de oportunidades buscadas", "Empleo, prácticas, becas, concursos, consultoría u otras", 2)

doc.add_heading("2. Resumen y propuesta de valor", level=1)
field("Resumen profesional", "60 a 90 palabras; indique intereses, métodos y el valor que puede aportar", 4)
field("Tres fortalezas verificables")

page_break()
doc.add_heading("3. Enlaces profesionales", level=1)
note("Compruebe que cada enlace sea público y abra sin solicitar permisos. No entregue contraseñas ni tokens de acceso.")
repeated_table(["Plataforma", "URL completa", "Qué encontrará el visitante"], [1900, 4100, 3360], 6)
field("LinkedIn")
field("GitHub")
field("ORCID / Google Scholar", "Si aplica")
field("Portafolio o sitio previo", "Si aplica")
field("Repositorio de datos", "Zenodo, OSF, Kaggle u otro")

doc.add_heading("4. Formación académica", level=1)
repeated_table(["Institución y programa", "Periodo", "Estado / distinción"], [4700, 1800, 2860], 3)
field("Tema de titulación o línea de interés", "Si todavía no existe, escriba PENDIENTE", 2)

page_break()
doc.add_heading("5. Proyectos", level=1)
note("Registre un proyecto por bloque. Describa su aporte real; no confunda una tarea de clase con experiencia laboral.")
for number in range(1, 4):
    doc.add_heading(f"Proyecto {number}", level=2)
    repeated_table(["Campo", "Información"], [2300, 7060], 5)
    p = doc.paragraphs[-1]
    labels = ["Título", "Problema u objetivo", "Datos y fuente", "Métodos y herramientas", "Resultado verificable"]
    table = doc.tables[-1]
    for idx, label in enumerate(labels, start=1):
        table.cell(idx, 0).text = label
        for run in table.cell(idx, 0).paragraphs[0].runs:
            set_font(run, 8.5, True, GREEN)
    field(f"Enlaces del proyecto {number}", "GitHub, demo, informe PDF, presentación, dataset o video")
    if number < 3:
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

page_break()
doc.add_heading("6. Publicaciones y productos académicos", level=1)
note("Use el estado exacto: publicado, aceptado, presentado, en revisión, manuscrito, working paper o trabajo de clase. Una buena etiqueta evita vender humo académico, especie sorprendentemente abundante.")
repeated_table(["Referencia completa", "Estado", "DOI / URL / documento"], [4500, 1500, 3360], 5)

doc.add_heading("7. Experiencia y participación", level=1)
repeated_table(["Organización y rol", "Fechas", "Responsabilidades y resultados"], [2900, 1500, 4960], 4)
field("Enlaces de evidencia", "Certificados, cartas, páginas institucionales o productos")

doc.add_heading("8. Cursos y certificaciones", level=1)
repeated_table(["Curso e institución", "Fecha / duración", "Credencial o URL"], [3900, 2100, 3360], 5)

page_break()
doc.add_heading("9. Competencias", level=1)
repeated_table(["Competencia", "Nivel y evidencia", "Proyecto donde se aplicó"], [2500, 3200, 3660], 6)
field("Software y lenguajes", "Python, R, Stata, Excel, SQL, Power BI u otros")
field("Métodos", "Econometría, series de tiempo, evaluación de impacto, SIG u otros")
field("Idiomas", "Indique nivel y certificación si existe")

doc.add_heading("10. Marca personal", level=1)
field("Tres adjetivos para la marca")
field("Colores preferidos y colores prohibidos")
field("Referencias visuales", "Enlaces a 2 o 3 sitios cuyo estilo le agrade")
field("Fotografía", "Enlace a Drive/GitHub con permiso público o indique que la adjuntará")
field("Biografía corta", "Máximo 30 palabras", 2)

page_break()
doc.add_heading("11. Archivos que debe entregar", level=1)
check_items = [
    "Hoja de vida actual, aunque esté incompleta.",
    "Fotografía profesional autorizada en buena resolución.",
    "Certificados que desea mostrar.",
    "Informes, papers, presentaciones o pósteres seleccionados.",
    "Enlaces públicos a repositorios y demostraciones.",
    "Logotipos institucionales solo si cuenta con autorización de uso.",
]
for text in check_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("☐  ")
    set_font(r, 12, False, GREEN)
    r = p.add_run(text)
    set_font(r, 10.5, False, INK)

doc.add_heading("12. Autorización y verificación", level=1)
for statement in [
    "Confirmo que la información entregada es verdadera y puede ser verificada.",
    "Autorizo la publicación del correo y los enlaces indicados como públicos.",
    "Entiendo que la fotografía generada de demostración se reemplazará antes de usar el portafolio profesionalmente.",
    "Revisaré la web y el PDF antes de su publicación definitiva.",
]:
    p = doc.add_paragraph()
    r = p.add_run("☐  ")
    set_font(r, 12, False, GREEN)
    r = p.add_run(statement)
    set_font(r, 10.5, False, INK)

field("Nombre de quien entrega la información")
field("Fecha")
field("Observaciones finales", "Datos que no deben publicarse o instrucciones especiales", 3)

doc.add_heading("13. Instrucción para entregar este archivo a la IA", level=1)
note("Adjunte este documento completo y escriba: 'Sustituye solo los datos de demostración. No inventes información. Señala contradicciones y campos pendientes. Actualiza la web, regenera la hoja de vida PDF y conserva una lista de verificación antes de publicar'.")

doc.core_properties.title = "Formulario de datos para portafolio profesional de Economía"
doc.core_properties.subject = "Recopilación de información verificable para web y hoja de vida"
doc.core_properties.author = "Proyecto Portafolio Profesional"
doc.save(OUT)
PUBLIC.write_bytes(OUT.read_bytes())
print(f"Generated: {OUT}")
print(f"Published: {PUBLIC}")
